from copy import deepcopy
from pathlib import Path
import re
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "baseline" / "v2.0" / "full-source" / "policies"
ZIP_OUT = ROOT / "baseline" / "v2.0" / "full-source" / "MotoCo_Policies_v2.0_Draft.zip"

POLICY_ZIPS = [
    ROOT / "policy 1-7.zip",
    ROOT / "policy8-14.zip",
    ROOT / "policy 11-22.zip",
    ROOT / "policy 23-27.zip",
]

GENERAL_NOTES = [
    "Recreated as part of Moto and Co Couriers baseline v2.0.",
    "Status is Draft for approval unless a signed approval reference is later added.",
    "Unconfirmed legal, owner, privacy, platform, or evidence details remain marked as TBD.",
]

POLICY_NOTES = {
    "Policy-03-PrivacyPolicy.docx": [
        "Add portal handling for customer, booking, delivery, receiver name/signature, billing, audit, and proof data.",
        "Add same-device offline outbox/local cache handling under GM Moto & Co Logistics Privacy Owner governance.",
        "Supabase data-location evidence remains required before final public privacy wording.",
    ],
    "Policy-04-CollectionNotice.docx": [
        "Add receiver name/signature, customer registration, booking, and offline local-device processing notices.",
        "APP 5 consequence-of-non-collection wording remains TBD.",
    ],
    "Policy-05-DataRetentionDestruction.docx": [
        "Add local device cache/outbox retention and destruction rule.",
        "POD proof retention remains 7 years from delivery date.",
        "Privacy Owner is role-based GM Moto & Co Logistics; legal-hold handling remains TBD.",
    ],
    "Policy-06-NDBResponsePlan.docx": [
        "Add lost/stolen driver device and unsynced local outbox as suspected incident triggers.",
        "Privacy Owner is role-based GM Moto & Co Logistics; OAIC notification artefacts, affected-person template, public statement process, and Digiverse handoff remain TBD.",
    ],
    "Policy-07-InformationSecurity.docx": [
        "Add supported device/browser, offline local cache, retry sync, clear-device-data risk, RLS/Storage UAT, and production access logging.",
        "Supported driver device/browser is approved for UAT evidence and must be re-executed after ERD approval and schema/runtime reconciliation.",
    ],
    "Policy-08-FailedDelivery.docx": [
        "Confirm offline retry/manual recovery impact on failed-delivery evidence and redelivery-fee billing.",
        "Final failed-delivery taxonomy and production accounting treatment remain TBD.",
    ],
    "Policy-09-PricingSchedule.docx": [
        "Align with corrected standard freight pricing and table-driven price_rules.",
        "Drivers must not manually enter prices.",
        "Production authority model remains TBD where Digiverse execution is required after Admin/Owner approval.",
    ],
    "Policy-10-InvoicingPaymentTerms.docx": [
        "Replace external accounting dependency with V1 portal-generated downloadable invoice PDF only.",
        "Admin email, bounce handling, payment follow-up, bank reconciliation, BAS/accountant handoff, OpenClaw, Xero, and accounting API integration are outside the V1 runtime.",
    ],
    "Policy-10a-CreditControl.docx": [
        "Align overdue notices, payment evidence, suspension/reinstatement, and manual invoice dispatch to V1 manual billing.",
        "Payment follow-up and bank reconciliation are off-system human/accounting processes; repeated non-payment termination pathway remains TBD.",
    ],
    "Policy-11-CustomerTerms.docx": [
        "Add portal booking, scheduled run status, next-run/brought-forward handling, POD download, disputes, manual invoice PDF, and no-SLA-monitoring boundary.",
        "Customer-facing legal copy remains TBD until approved.",
    ],
    "Policy-12-DriverAgreement.docx": [
        "Held outside logistics portal scope and preserved for future HCM system.",
    ],
    "Policy-13-DriverVerification.docx": [
        "Held outside logistics portal scope and preserved for future HCM system.",
    ],
    "Policy-14-OrderCancellation.docx": [
        "Confirm customer-visible cancellation/review wording and manual notification evidence.",
        "External notification channel remains TBD.",
    ],
    "Policy-15-GoodsAcceptance.docx": [
        "Align with pickup-time item counting and driver dock-decision evidence.",
        "Policy #15 remains Draft/Awaiting Review until owner approval is provided.",
    ],
    "Policy-16-VendorPickupStandards.docx": [
        "Align active supplier list and supplier archive/reactivation evidence.",
        "Ficeda is removed from active V1 supplier network.",
        "Named dock contacts, supplier review cadence, and health scoring remain TBD.",
    ],
    "Policy-18-DeliveryDispute.docx": [
        "Add POD download, manual invoice PDF, billing-query acknowledgement, remedy, and credit-note/corrected-invoice path.",
        "Invoice emailing and bank reconciliation remain outside runtime; actual credit-note/corrected-invoice issue/send path remains TBD if later required.",
    ],
    "Policy-19-DriverCodeOfConduct.docx": [
        "Held outside logistics portal scope and preserved for future HCM system.",
    ],
    "Policy-20-AIUsePolicy.docx": [
        "Keep live AI generation/send out of V1 unless separately approved.",
        "Local AI draft review may exist; approved drafts are not sent.",
    ],
    "Policy-21-InternalAcceptableUseData.docx": [
        "Add production access logging, offline local device data handling, export approval evidence, and RLS/Auth boundary UAT.",
        "Digiverse data-processing/security schedule and access-log format remain TBD.",
    ],
    "Policy-22-DriverScheduling.docx": [
        "Keep logistics availability only; do not reintroduce HCM/legal classification.",
        "Owner could not locate Policy #22 source for approval; previous evening-before lockdown is rejected.",
        "New runtime requirement: driver creates daily run from ready con notes and can collect planned, bring-forward, and no-con-note depot packages.",
    ],
    "Policy-23-AccountSuspensionTermination.docx": [
        "Keep repeated non-payment termination blocked.",
        "Align manual invoice and notice evidence.",
        "Debt recovery escalation path, write-off thresholds, authority, and notice wording remain TBD.",
    ],
    "Policy-24-RevenueReportingFinancialControls.docx": [
        "Add downloadable EOM invoice PDF as the V1 runtime revenue artifact.",
        "Admin email, payment follow-up, bank reconciliation, BAS/accountant handoff, and any corrected-invoice handling are outside V1 runtime.",
        "No Xero, OpenClaw, or accounting API integration is part of the V1 baseline.",
    ],
    "Policy-25-PickupDeliveryExecutionMultiDriver.docx": [
        "Conditional future activation only; not active logistics portal scope for V1.",
    ],
    "Policy-26-BillingPaymentCouriers.docx": [
        "Conditional future external courier/payment model only; not active logistics portal scope for V1.",
    ],
    "Policy-27-WHSFatigueDriverWellbeing.docx": [
        "Keep active supplier WHS hazard and No Pickup evidence.",
        "Full fatigue/risk framework remains conditional pending trigger and owner/process approval.",
    ],
}

POLICY_CURRENT_RULES = {
    "Policy-03-PrivacyPolicy.docx": [
        "Portal data includes customer/profile, booking, delivery, receiver name/signature, invoice/PDF, audit, and proof records.",
        "Driver offline mode may hold unsynced delivery updates locally on the signed-in device until sync or Admin recovery succeeds, then clears.",
        "Privacy Owner is role-based GM Moto & Co Logistics.",
        "Australian data residency evidence remains required before final public wording.",
    ],
    "Policy-04-CollectionNotice.docx": [
        "Collection notice must cover customer registration, pickup requests, receiver name, receiver signature, POD proof, billing records, and offline local-device processing.",
        "Consequences of not collecting mandatory receiver name/signature remain TBD for final legal wording.",
    ],
    "Policy-05-DataRetentionDestruction.docx": [
        "POD proof is retained for 7 years from delivery date.",
        "Unsynced local outbox/cache records must be cleared after confirmed sync or Admin-approved recovery.",
        "Privacy Owner is role-based GM Moto & Co Logistics; legal-hold handling remains TBD.",
    ],
    "Policy-06-NDBResponsePlan.docx": [
        "Lost/stolen driver device, exposed local cache, failed proof sync, or unauthorized access to POD data are suspected incident triggers.",
        "Privacy Owner is role-based GM Moto & Co Logistics.",
        "OAIC notice artefacts, affected-person template, public statement process, and Digiverse handoff remain TBD.",
    ],
    "Policy-07-InformationSecurity.docx": [
        "Driver offline mode is allowed only for retry sync and recovery; it is not a second source of truth.",
        "Supported driver device/browser must be captured in UAT evidence and re-executed after ERD approval and schema/runtime reconciliation.",
        "Production RLS proof, Storage proof, and access log evidence remain required.",
    ],
    "Policy-08-FailedDelivery.docx": [
        "Failed-delivery evidence must survive offline retry and Admin recovery.",
        "Final failed-delivery taxonomy and accounting treatment remain TBD.",
    ],
    "Policy-09-PricingSchedule.docx": [
        "Current standard freight pricing is exclusive of GST.",
        "Tyres: 1 tyre $18.50; 2 tyres $24.00; 3 tyres $33.00; 4+ tyres $12.30 each.",
        "General parts: up to 5kg $17.20; 5-10kg $21.00; 10kg+ from $25.00 subject to handling approval.",
        "Additional charges: return to supplier pre-labelled $6.00; out-of-zone delivery $10.00; oversized/bulky freight quoted on application.",
        "Pricing must be table-driven from price_rules and must not be manually entered by a driver.",
    ],
    "Policy-10-InvoicingPaymentTerms.docx": [
        "V1 invoices are generated as portal PDF downloads by Admin.",
        "Admin emails clients manually outside the portal.",
        "No accounting API/OpenClaw/Xero integration is part of the V1 baseline.",
        "Email bounce handling, payment follow-up, bank reconciliation, and BAS/accountant handoff are off-system human/accounting processes.",
    ],
    "Policy-10a-CreditControl.docx": [
        "Credit control must rely on generated invoice PDF evidence and Admin review inside the portal.",
        "Email/payment/bank reconciliation steps are off-system human/accounting processes.",
        "Repeated non-payment termination, write-off authority, and debt recovery escalation remain TBD.",
    ],
    "Policy-11-CustomerTerms.docx": [
        "Customers can register, be approved by Admin, create pickup requests, view scheduled/current/delivered orders, download POD, and raise review/dispute requests.",
        "Scheduled run status and brought-forward handling must be visible; the logistics portal does not perform SLA monitoring.",
        "Customer-facing legal copy remains TBD until approved.",
    ],
    "Policy-12-DriverAgreement.docx": [
        "This policy is HCM scope and is not controlling logistics portal runtime for V1.",
    ],
    "Policy-13-DriverVerification.docx": [
        "This policy is HCM scope and is not controlling logistics portal runtime for V1.",
    ],
    "Policy-14-OrderCancellation.docx": [
        "Customer-visible cancellation and review wording must match the portal; external notification channel remains TBD.",
    ],
    "Policy-15-GoodsAcceptance.docx": [
        "Driver item count occurs at pickup before goods are moved to en route/sign-off.",
        "Receiver phone is not mandatory for delivered status; receiver name and signature are mandatory.",
    ],
    "Policy-16-VendorPickupStandards.docx": [
        "Active V1 suppliers are Link International, A1 Accessories, McLeods, Gas Imports, and Whites Powersports.",
        "Ficeda is removed from the active supplier network and must require approval/evidence before reactivation.",
        "Named dock contacts, supplier review cadence, and health scoring remain TBD.",
    ],
    "Policy-18-DeliveryDispute.docx": [
        "Client review/dispute actions must be recorded, confirmed to the client, and visible for Admin investigation.",
        "Admin investigates against pickup count, POD PDF/signature, invoice PDF, and audit history.",
        "Invoice email and bank reconciliation are outside runtime; credit-note/corrected-invoice issue/send path remains TBD if later required.",
    ],
    "Policy-19-DriverCodeOfConduct.docx": [
        "This policy is HCM scope and is not controlling logistics portal runtime for V1.",
    ],
    "Policy-20-AIUsePolicy.docx": [
        "AI draft generation may exist only as reviewed internal drafting support; AI-generated sends are not active V1 behaviour.",
    ],
    "Policy-21-InternalAcceptableUseData.docx": [
        "Production access, offline local data, export/download actions, POD data, and role changes require audit evidence.",
        "Digiverse data-processing/security schedule and access-log format remain TBD.",
    ],
    "Policy-22-DriverScheduling.docx": [
        "Driver schedule/runs are logistics availability only; HCM/legal classification is not determined by the portal.",
        "Owner could not locate Policy #22 source for approval; source review remains open.",
        "Previous evening-before lockdown is rejected.",
        "Driver must be able to create a daily run from ready con notes.",
        "At each depot, driver can collect planned milk-run package, bring forward ready next-day package, or record ready package with no con note/customer missed portal entry.",
    ],
    "Policy-23-AccountSuspensionTermination.docx": [
        "Account suspension/reinstatement must align to manual invoice and notice evidence.",
        "Repeated non-payment termination remains blocked until legal/owner approval.",
    ],
    "Policy-24-RevenueReportingFinancialControls.docx": [
        "Runtime revenue evidence consists of completed delivery/POD and generated invoice PDF.",
        "Admin email, payment follow-up, bank reconciliation, BAS/accountant handoff, and any corrected-invoice handling are outside V1 runtime.",
        "No Xero, OpenClaw, or accounting API integration is part of the V1 baseline.",
        "Otimi Rules reporting remains outside runtime if required.",
    ],
    "Policy-25-PickupDeliveryExecutionMultiDriver.docx": [
        "Conditional future activation only; not active V1 logistics portal scope.",
    ],
    "Policy-26-BillingPaymentCouriers.docx": [
        "Conditional future external courier/payment model only; not active V1 logistics portal scope.",
    ],
    "Policy-27-WHSFatigueDriverWellbeing.docx": [
        "Active supplier WHS hazard/no-pickup evidence remains in scope; broader fatigue/risk framework remains conditional pending owner/process approval.",
    ],
}

POLICY_BODY_REWRITES = {
    "Policy-16-VendorPickupStandards.docx": {
        "2.  The milk run is compiled the night before by APP-ADM-002 (Run Planning Module) via pg_cron. Suppliers are not notified on the day of the run — goods must be ready by 10:00am regardless of the driver's estimated arrival time.": "2.  V1 does not use a hard night-before run lockdown. Supplier/customer con-note timing is an upstream input, so the driver creates the daily run from ready con notes. Suppliers are not notified on the day of the run; goods should be ready by 10:00am where supplier process allows, and any not-ready goods are handled through the approved No Pickup or depot collection evidence path.",
    },
    "Policy-22-DriverScheduling.docx": {
        "1.  Drivers must notify Admin of any unavailability (leave, illness, personal commitment) as early as possible and no later than the evening before a scheduled run.": "1.  Drivers must notify Admin of any unavailability (leave, illness, personal commitment) as early as possible. V1 keeps light-touch logistics availability governance and does not use the previous evening-before run lockdown as the operating control.",
        "2.  A driver who cannot perform a scheduled run must notify Admin before APP-ADM-002 compiles the run (before the pg_cron trigger runs the night before).": "2.  A driver who cannot perform a run must notify Admin before departure planning is relied on for that run. APP-ADM-002/driver daily-run evidence must show whether the named driver is available before the run departs.",
        "4.  APP-ADM-002 assigns the available driver and a named vehicle to each run the night before. A run must not depart without both a named driver and vehicle assigned (POL-MCL-002-001).": "4.  V1 requires a named driver and named vehicle before departure. The driver-created daily run consolidates ready con notes, and a run must not depart without both a named driver and vehicle assigned (POL-MCL-002-001).",
    },
    "Policy-10-InvoicingPaymentTerms.docx": {
        "Generates invoice on Admin approval; dispatches to ACT-CRM-001b": "Generates downloadable invoice PDF on Admin approval; Admin emails the PDF manually outside the portal",
    },
}


def versioned_name(original_name):
    stem = original_name[:-5] if original_name.lower().endswith(".docx") else original_name
    return f"{stem}-v2.0.docx"


def insert_front_paragraph(doc, text, *, bold=False, size=10, color=None, align=None):
    body = doc.element.body
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    body.remove(p._p)
    body.insert(0, p._p)


def add_front_matter(doc, original_name):
    notes = POLICY_NOTES.get(original_name, ["Reviewed against baseline v2.0. No specific operating rule change identified."])
    current_rules = POLICY_CURRENT_RULES.get(
        original_name,
        ["No specific v2.0 runtime rule change identified. Preserved source content remains subject to owner/legal approval."],
    )
    lines = []
    lines.append("Moto and Co Couriers Policy Baseline v2.0")
    lines.append(f"Source file: {original_name}")
    lines.append("Version: v2.0")
    lines.append("Status: Draft for approval - not final legal copy")
    lines.append("Created: 2026-07-02")
    lines.append("Owner: TBD")
    lines.append("Legal approval reference: TBD")
    lines.append("Baseline pack: baseline/v2.0/full-source/policies")
    lines.append("")
    lines.append("Version control note:")
    for note in GENERAL_NOTES + notes:
        lines.append(f"- {note}")
    lines.append("")
    lines.append("Current v2.0 baseline rules:")
    lines.append("The following rules supersede conflicting preserved v1.0 text unless a later signed owner/legal approval says otherwise.")
    for rule in current_rules:
        lines.append(f"- {rule}")
    lines.append("")
    lines.append("Preserved v1 source content follows for traceability. If it conflicts with the v2.0 baseline rules above, the v2.0 rule controls while this document remains Draft for approval.")
    lines.append("")

    for line in reversed(lines):
        if line.startswith("- "):
            insert_front_paragraph(doc, line, size=9)
        elif line == "Moto and Co Couriers Policy Baseline v2.0":
            insert_front_paragraph(doc, line, bold=True, size=16, color="E11D48", align=WD_ALIGN_PARAGRAPH.LEFT)
        elif line.endswith(":"):
            insert_front_paragraph(doc, line, bold=True, size=10)
        else:
            insert_front_paragraph(doc, line, size=9)

def replace_paragraph_text(paragraph, new_text):
    for run in list(paragraph.runs):
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def apply_body_rewrites(doc, original_name):
    rewrites = POLICY_BODY_REWRITES.get(original_name, {})
    if not rewrites:
        return
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if text in rewrites:
            replace_paragraph_text(paragraph, rewrites[text])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    if text in rewrites:
                        replace_paragraph_text(paragraph, rewrites[text])


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    manifest = [["policy_file", "versioned_file", "version", "status", "source_archive", "change_note"]]
    created = []
    for archive_path in POLICY_ZIPS:
        with zipfile.ZipFile(archive_path) as archive:
            for member in archive.namelist():
                if not member.lower().endswith(".docx") or member.startswith("__MACOSX/"):
                    continue
                original_name = Path(member).name
                temp_path = OUT / f"__tmp_{original_name}"
                temp_path.write_bytes(archive.read(member))
                doc = Document(temp_path)
                apply_body_rewrites(doc, original_name)
                add_front_matter(doc, original_name)
                doc.core_properties.title = f"{original_name} v2.0"
                doc.core_properties.subject = "Moto and Co Couriers policy baseline v2.0"
                doc.core_properties.author = "Moto and Co Couriers"
                output_name = versioned_name(original_name)
                output_path = OUT / output_name
                doc.save(output_path)
                temp_path.unlink(missing_ok=True)
                notes = POLICY_NOTES.get(original_name, ["Reviewed against baseline v2.0."])
                manifest.append([
                    original_name,
                    output_name,
                    "v2.0",
                    "Draft for approval",
                    archive_path.name,
                    " | ".join(notes),
                ])
                created.append(output_path)

    manifest_path = OUT / "Policy_v2_baseline_manifest.csv"
    manifest_path.write_text(
        "\n".join(",".join(f'"{str(cell).replace(chr(34), chr(34) + chr(34))}"' for cell in row) for row in manifest),
        encoding="utf-8",
    )

    ZIP_OUT.unlink(missing_ok=True)
    with zipfile.ZipFile(ZIP_OUT, "w", zipfile.ZIP_DEFLATED) as zip_out:
        for path in created + [manifest_path]:
            zip_out.write(path, arcname=path.name)

    print(f"Created {len(created)} policy docs")
    print(f"Created {ZIP_OUT}")


if __name__ == "__main__":
    main()
