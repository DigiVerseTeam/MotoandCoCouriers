from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BASELINE_DIR = ROOT / "baseline" / "v2.0"
RENDER_DIR = BASELINE_DIR / "_rendered"


DOCS = [
    ("MotoCo_Baseline_Documentation_Control_v2.0.md", "MotoCo_Baseline_Documentation_Control_v2.0.docx"),
    ("MotoCo_BOAS_Baseline_Addendum_v2.0.md", "MotoCo_BOAS_Baseline_Addendum_v2.0.docx"),
    ("MotoCo_SOP_Baseline_Addendum_v1.3.md", "MotoCo_SOP_Baseline_Addendum_v1.3.docx"),
    ("MotoCo_Policy_Baseline_Addendum_v2.0.md", "MotoCo_Policy_Baseline_Addendum_v2.0.docx"),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B9B9C3", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_margins(section):
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def add_header_footer(doc, title):
    section = doc.sections[0]
    header = section.header
    if header.paragraphs:
        p = header.paragraphs[0]
    else:
        p = header.add_paragraph()
    p.text = "Moto and Co Couriers baseline pack"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(85, 85, 85)

    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = f"{title} | Draft for approval"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(85, 85, 85)


def parse_table(lines):
    rows = []
    for line in lines:
        parts = [part.strip() for part in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", part or "") for part in parts):
            continue
        rows.append(parts)
    return rows


def column_widths(col_count):
    patterns = {
        2: [2700, 6660],
        3: [2200, 3800, 3360],
        4: [1900, 3100, 2600, 1760],
    }
    return patterns.get(col_count, [int(9360 / col_count)] * col_count)


def add_table(doc, table_lines):
    rows = parse_table(table_lines)
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_borders(table)
    widths = column_widths(col_count)
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            cell = table.cell(r_idx, c_idx)
            text = row[c_idx] if c_idx < len(row) else ""
            cell.text = text
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, widths[c_idx])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9 if col_count >= 4 else 10)
        if r_idx == 0:
            set_repeat_table_header(table.rows[r_idx])
            for cell in table.rows[r_idx].cells:
                set_cell_shading(cell, "E8EEF5")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
    doc.add_paragraph()


def add_title(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(225, 29, 72)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("Version-controlled draft for approval")
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(85, 85, 85)


def add_markdown(doc, markdown_text):
    lines = markdown_text.splitlines()
    table_buffer = []
    title_seen = False

    def flush_table():
        nonlocal table_buffer
        if table_buffer:
            add_table(doc, table_buffer)
            table_buffer = []

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            flush_table()
            continue
        if line.startswith("|"):
            table_buffer.append(line)
            continue
        flush_table()

        if line.startswith("# "):
            title = line[2:].strip()
            if not title_seen:
                add_title(doc, title)
                title_seen = True
            else:
                doc.add_heading(title, level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            p.add_run(line[2:].strip())
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(4)
            p.add_run(re.sub(r"^\d+\. ", "", line).strip())
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.add_run(line)
    flush_table()


def build_docx(md_path, docx_path):
    markdown_text = md_path.read_text(encoding="utf-8")
    title = markdown_text.splitlines()[0].lstrip("# ").strip()
    doc = Document()
    for section in doc.sections:
        set_margins(section)
    configure_styles(doc)
    add_header_footer(doc, title)
    add_markdown(doc, markdown_text)
    doc.core_properties.title = title
    doc.core_properties.subject = "Moto and Co Couriers baseline documentation"
    doc.core_properties.author = "Moto and Co Couriers"
    doc.save(docx_path)


def main():
    BASELINE_DIR.mkdir(parents=True, exist_ok=True)
    for md_name, docx_name in DOCS:
        build_docx(BASELINE_DIR / md_name, BASELINE_DIR / docx_name)
        print(f"created {BASELINE_DIR / docx_name}")


if __name__ == "__main__":
    main()
