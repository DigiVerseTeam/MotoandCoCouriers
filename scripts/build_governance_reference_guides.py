from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\User\OneDrive\Documents\Moto and Co Couriers")
OUT_DIR = ROOT / "Governance Framework"
TODAY = "23 July 2026"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(85, 85, 85)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)


def set_paragraph_font(paragraph, size=11, color=None, bold=None, italic=None):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)


def configure_document(doc, title):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for list_style in ("List Bullet", "List Number"):
        style = styles[list_style]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = f"Moto & Co Couriers | Governance Framework"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_font(header, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    add_page_number(footer)


def add_title_block(doc, title, subtitle, guide_id):
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(2)
    r = kicker.add_run("MOTO & CO COURIERS")
    set_run_font(r, size=10, color=MUTED, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=24, color=INK, bold=True)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(14)
    r = sub.add_run(subtitle)
    set_run_font(r, size=12.5, color=MUTED)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [1800, 7560])
    rows = [
        ("Guide ID", guide_id),
        ("Status", "Working reference guide"),
        ("Created", TODAY),
        ("Use", "Internal governance and documentation alignment"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                set_paragraph_font(p, size=10)
            for run in cell.paragraphs[0].runs:
                if cell is row.cells[0]:
                    run.bold = True
    doc.add_paragraph()


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    set_run_font(r, size=11, color=INK, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_label_table(doc, rows, widths=(2300, 7060), header=None):
    extra = 1 if header else 0
    table = doc.add_table(rows=len(rows) + extra, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, list(widths))
    row_offset = 0
    if header:
        table.cell(0, 0).text = header[0]
        table.cell(0, 1).text = header[1]
        for cell in table.rows[0].cells:
            set_cell_shading(cell, LIGHT_BLUE)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                set_paragraph_font(p, size=10, bold=True)
        row_offset = 1
    for idx, (label, value) in enumerate(rows, row_offset):
        table.cell(idx, 0).text = label
        table.cell(idx, 1).text = value
        set_cell_shading(table.cell(idx, 0), LIGHT_GRAY)
        for cell in table.rows[idx].cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                set_paragraph_font(p, size=10)
        for run in table.cell(idx, 0).paragraphs[0].runs:
            run.bold = True
    doc.add_paragraph()


def add_matrix_table(doc, headers, rows, widths):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = header
        set_cell_shading(cell, LIGHT_BLUE)
    for r_idx, row in enumerate(rows, 1):
        for c_idx, value in enumerate(row):
            table.cell(r_idx, c_idx).text = value
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                set_paragraph_font(p, size=9.5)
    for cell in table.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    doc.add_paragraph()


def add_relationship_map(doc, current):
    rows = [
        ("Capability Charter", "Why the capability exists, what it must become, and what structural rules govern it.", "Strategic intent and enduring authority."),
        ("Operational Policy", "What rule must be followed in operations, including scope, roles, exceptions, and review.", "Standing operational rule."),
        ("SOP", "How the work is performed, evidenced, checked, and escalated.", "Executable process."),
    ]
    marked = []
    for name, purpose, role in rows:
        if name == current:
            name = f"{name} (this document type)"
        marked.append((name, purpose, role))
    add_matrix_table(
        doc,
        ["Document type", "Question it answers", "Authority level"],
        marked,
        [2300, 4760, 2300],
    )


def source_note(doc, items):
    doc.add_heading("Source Alignment", level=1)
    add_bullets(doc, items)


def save_doc(doc, filename):
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUT_DIR / filename
    doc.save(path)
    return path


def build_capability_guide():
    doc = Document()
    configure_document(doc, "Capability Charter Reference Guide")
    add_title_block(
        doc,
        "Capability Charter Reference Guide",
        "How Moto & Co defines capability-level policy, strategic intent, and operational linkage.",
        "GOV-REF-001",
    )
    add_callout(
        doc,
        "Plain-English definition",
        "A Capability Charter is the durable governance document for a business capability: it explains why the capability exists, what outcome it must deliver, what strategic posture it holds, and which enduring capability policies constrain how it operates.",
    )
    add_relationship_map(doc, "Capability Charter")

    doc.add_heading("What It Is", level=1)
    add_bullets(
        doc,
        [
            "The top-level operating authority for a named capability such as Network & Supplier Access or Run Planning & Dispatch.",
            "A home for durable strategic intent: posture, horizon, investment philosophy, maturity target, deployment stance, risks, and evolution outlook.",
            "A home for capability policies: standing rules that constrain the capability even if the operating process changes.",
            "The bridge between the business model and the operating layer: BOAS, operational policies, SOPs, roles, data, risks, and open gaps.",
        ],
    )

    doc.add_heading("What It Is Not", level=1)
    add_bullets(
        doc,
        [
            "It is not a step-by-step work instruction.",
            "It is not a legal/public policy document.",
            "It is not a temporary project plan or implementation backlog.",
            "It should not duplicate every SOP, workflow screen, or platform field.",
        ],
    )

    doc.add_heading("Decision Tests", level=1)
    add_label_table(
        doc,
        [
            ("Redesign Test", "A capability policy belongs here if it would still hold if the operating process were redesigned tomorrow."),
            ("Reorganisation Test", "A charter statement belongs here if it would still hold if the capability were delivered through a different team, system, or partner model."),
            ("Durability Test", "If the statement is a short-term implementation action, put it in a backlog, addendum, or SOP change register instead."),
            ("Authority Test", "If a downstream policy or SOP conflicts with the capability charter, the capability-level authority wins until formally amended."),
        ],
    )

    doc.add_heading("Required Outline", level=1)
    add_numbers(
        doc,
        [
            "Capability identity: ID, name, role, strategic posture, owner, maturity, deployment, lifecycle state, review cadence.",
            "Capability outcome: the result the business must be able to produce repeatedly.",
            "Value-chain position: where the capability sits in the service path and what breaks if it fails.",
            "Business value anchors: revenue, customer outcome, goodwill, resilience, compliance, or other weighted anchors.",
            "L2 sub-capabilities: the smaller capability components that make the parent capability work.",
            "Capability policy: policy ID, policy statement, owner, source, affected capabilities, redesign test, review cadence.",
            "Capability charter: charter ID, horizon, strategic posture, directional intent, investment philosophy, deployment stance, value trajectory, maturity target, risks, evolution outlook, decision rights, measures, and review triggers.",
            "Operational links: BOAS references, governed operational policies, relevant SOPs, gaps, and document control.",
        ],
    )

    doc.add_heading("How It Connects Downstream", level=1)
    add_matrix_table(
        doc,
        ["Layer", "Capability charter decides", "Downstream document does"],
        [
            ("Operational policy", "The standing rule or boundary that must hold.", "Turns the rule into operational obligations, roles, exceptions, and review cadence."),
            ("SOP", "The capability outcome and non-negotiable constraints.", "Turns the policy into repeatable actions, triggers, evidence, runtime mapping, and escalation."),
            ("BOAS", "The capability identity and operating architecture.", "Registers the process, roles, access, data objects, controls, and system references."),
        ],
        [2100, 3600, 3660],
    )

    doc.add_heading("Moto & Co Examples", level=1)
    add_bullets(
        doc,
        [
            "CAP-MCL-001 Network & Supplier Access uses a Differentiator posture because supplier relationships and dock access are a competitive asset.",
            "CAP-MCL-002 Run Planning & Dispatch uses a Compete posture because reliable planning is essential, but clients ultimately value delivery outcomes.",
            "The capability policy examples include the Supplier Approval Gate and the Night-Before Compilation & Named Assignment rule.",
            "Each capability document links back to BOAS implementation, operational policies, open gaps, and document control.",
        ],
    )

    doc.add_heading("Quality Checklist", level=1)
    add_bullets(
        doc,
        [
            "The capability outcome is clear enough that a reader can tell what business ability is being governed.",
            "The charter has a real horizon and strategic posture, not only current operating detail.",
            "Capability policies pass the Redesign Test.",
            "Operational policies and SOPs are linked, not pasted wholesale into the charter.",
            "Open gaps are named with impact and next review trigger.",
        ],
    )

    source_note(
        doc,
        [
            "CAP-MCL-001-NetworkSupplierAccess.docx",
            "CAP-MCL-002-RunPlanningDispatch.docx",
            "baseline/v2.0/MotoCo_BOAS_Baseline_Addendum_v2.0.md",
            "baseline/v2.0/MotoCo_Baseline_Documentation_Control_v2.0.md",
        ],
    )
    return save_doc(doc, "MotoCo_Capability_Charter_Reference_Guide.docx")


def build_policy_guide():
    doc = Document()
    configure_document(doc, "Operational Policy Reference Guide")
    add_title_block(
        doc,
        "Operational Policy Reference Guide",
        "How Moto & Co turns capability authority into standing operating rules.",
        "GOV-REF-002",
    )
    add_callout(
        doc,
        "Plain-English definition",
        "An Operational Policy is a standing rule document: it states what must, must not, or may happen in operations, who is accountable, what exceptions exist, and which SOPs or systems must follow it.",
    )
    add_relationship_map(doc, "Operational Policy")

    doc.add_heading("What It Is", level=1)
    add_bullets(
        doc,
        [
            "A controlled rule set for a specific operational subject, such as invoicing, vendor pickup standards, failed delivery, pricing, privacy, or customer terms.",
            "The document that converts capability-level authority into enforceable operating obligations.",
            "The reference point for Admin, Driver, system modules, clients, suppliers, and other actors when decisions or exceptions occur.",
            "A stable enough rule source that multiple SOPs can implement it without rewriting the policy each time a screen or workflow changes.",
        ],
    )

    doc.add_heading("What It Is Not", level=1)
    add_bullets(
        doc,
        [
            "It is not a click-by-click procedure.",
            "It is not a system configuration file or migration script.",
            "It is not a vague principle without an accountable owner, scope, and exception path.",
            "It should not hide unresolved legal, privacy, owner, or evidence details; those stay marked as TBD until approved.",
        ],
    )

    doc.add_heading("When To Use An Operational Policy", level=1)
    add_bullets(
        doc,
        [
            "Use it when a business rule needs authority across people, systems, and SOPs.",
            "Use it when the rule affects risk, billing, privacy, safety, customer obligations, supplier obligations, or audit evidence.",
            "Use it when the organisation needs a formal source for exceptions and review cadence.",
            "Do not use it for one-off tasks, temporary project notes, or detailed workflow instructions.",
        ],
    )

    doc.add_heading("Required Outline", level=1)
    add_numbers(
        doc,
        [
            "Policy identity: policy number, policy ID, policy name, priority, status, parent capability policy, owner, effective date, review cadence, and lifecycle state.",
            "Purpose: why the rule exists and which risks it controls.",
            "Scope: who and what the policy applies to, including roles, customers, suppliers, systems, records, and service lines.",
            "Rules: numbered, testable rules that state what must happen, must not happen, or may happen only with approval.",
            "Roles and responsibilities: RACI or clear accountable owner, responsible roles, consulted parties, and informed parties.",
            "Linked processes: SOPs, BOAS references, system modules, data objects, and related policies.",
            "Exceptions: what exceptions are allowed, who approves them, and which exceptions are forbidden.",
            "Review cadence and document control: version, approval reference, owner, next review, and change triggers.",
        ],
    )

    doc.add_heading("Decision Tests", level=1)
    add_label_table(
        doc,
        [
            ("Rule Test", "If the statement says what must or must not happen, it probably belongs in a policy."),
            ("Procedure Test", "If the statement says the exact sequence of actions to perform, move that detail to an SOP."),
            ("Authority Test", "If people need permission, approval, accountability, or an exception path, the policy must say who holds that authority."),
            ("Evidence Test", "If compliance must be proven later, the policy should identify the required record or linked SOP evidence."),
        ],
    )

    doc.add_heading("Boundary With SOPs", level=1)
    add_matrix_table(
        doc,
        ["Question", "Policy answer", "SOP answer"],
        [
            ("Invoice approval", "No invoice is created without Admin approval.", "How Admin reviews the billing group and how the invoice is generated or downloaded."),
            ("Vendor pickup", "Goods must be labelled and ready; No Pickup cannot create a billable item.", "How the driver checks goods, records No Pickup, and captures evidence."),
            ("Data retention", "Which records must be retained and for how long.", "How retention queues, destruction checks, or recovery steps are executed."),
        ],
        [1800, 3780, 3780],
    )

    doc.add_heading("Moto & Co Examples", level=1)
    add_bullets(
        doc,
        [
            "Policy #10 Invoicing & Payment Terms governs the monthly invoice cycle, Admin approval gate, billing contact requirement, payment terms, and invoice content.",
            "Policy #16 Vendor Pickup Standards governs pickup windows, packaging, con note labelling, grace period, No Pickup handling, and bring-forward conditions.",
            "The v2.0 baseline keeps unconfirmed legal, owner, privacy, platform, or evidence details marked as TBD rather than burying them in procedure.",
            "Operational policies link upward to capability policies and downward to SOPs, BOAS records, and runtime evidence.",
        ],
    )

    doc.add_heading("Quality Checklist", level=1)
    add_bullets(
        doc,
        [
            "The policy has a clear parent capability or capability policy.",
            "Every rule is specific enough to test and broad enough to survive system-screen changes.",
            "Roles and exceptions are clear.",
            "Linked SOPs are named without duplicating their full procedures.",
            "Draft, legal, owner, privacy, or platform uncertainties remain visible as TBD until approved.",
        ],
    )

    source_note(
        doc,
        [
            "baseline/v2.0/full-source/policies/Policy-10-InvoicingPaymentTerms-v2.0.docx",
            "baseline/v2.0/full-source/policies/Policy-16-VendorPickupStandards-v2.0.docx",
            "baseline/v2.0/MotoCo_Policy_Baseline_Addendum_v2.0.md",
            "CAP-MCL-001-NetworkSupplierAccess.docx",
        ],
    )
    return save_doc(doc, "MotoCo_Operational_Policy_Reference_Guide.docx")


def build_sop_guide():
    doc = Document()
    configure_document(doc, "SOP Reference Guide")
    add_title_block(
        doc,
        "SOP Reference Guide",
        "How Moto & Co writes executable process documents that evidence policy compliance.",
        "GOV-REF-003",
    )
    add_callout(
        doc,
        "Plain-English definition",
        "A Standard Operating Procedure is the repeatable execution document: it tells the actor or system what happens, in what order, when it starts, what evidence proves completion, and what to do when the normal path fails.",
    )
    add_relationship_map(doc, "SOP")

    doc.add_heading("What It Is", level=1)
    add_bullets(
        doc,
        [
            "A controlled procedure for a specific process, event, or exception path.",
            "The execution layer beneath capability charters and operational policies.",
            "A training and audit tool: a reader should be able to perform, monitor, or test the process from the SOP.",
            "A bridge between human steps and runtime implementation: business steps stay platform-agnostic, while platform mapping, data, RACI, runtime, and SLA details sit in their own sections.",
        ],
    )

    doc.add_heading("What It Is Not", level=1)
    add_bullets(
        doc,
        [
            "It is not the policy itself and must not override policy.",
            "It is not the strategic charter for a capability.",
            "It is not only a software specification, even when the system performs some steps.",
            "It should not conceal policy changes; if a new rule appears during SOP drafting, update the policy or raise a controlled gap.",
        ],
    )

    doc.add_heading("Required Workbook Sections", level=1)
    add_matrix_table(
        doc,
        ["Section", "Purpose", "Must answer"],
        [
            ("Summary", "Identity, owner, parent policy, process module, purpose, outcomes, dependencies, completion standard.", "Why does this SOP exist and when is it complete?"),
            ("Steps", "Platform-agnostic process steps.", "Who does what, when, and what proves the step is complete?"),
            ("Skills", "Knowledge or capability needed to execute or review the process.", "What must the actor understand before doing this safely?"),
            ("Platform", "Technology mapping and BOAS/system references.", "Which app/module/table supports each step?"),
            ("Data & CIA", "Data objects, classification, owner, retention, and privacy notes.", "What records are created or changed and how sensitive are they?"),
            ("Runtime", "Workflow triggers, checkpoints, retry logic, success and failure conditions.", "How does the process run or fail in production?"),
            ("RACI", "Responsible, accountable, consulted, informed roles.", "Who owns the action and who must know?"),
            ("SLA", "Response and execution time obligations.", "How quickly must the action or system step happen?"),
        ],
        [1700, 4050, 3610],
    )

    doc.add_heading("Step Quality Rules", level=1)
    add_bullets(
        doc,
        [
            "Each step needs an actor, trigger, action, completion criteria, and exception or escalation path.",
            "Human steps and system steps must be separated clearly.",
            "The business step should stay readable even if the app screen changes.",
            "Every critical control needs evidence: approval record, timestamp, proof record, exception note, invoice ID, or similar.",
            "A failed step must say what happens next; silence is not an escalation path.",
        ],
    )

    doc.add_heading("Decision Tests", level=1)
    add_label_table(
        doc,
        [
            ("Execution Test", "If a trained person could follow it to perform or review the work, it belongs in an SOP."),
            ("Evidence Test", "If the business needs proof that the step happened, the SOP must name the record created or updated."),
            ("Policy Boundary Test", "If the text creates or changes a rule, it belongs in policy first, then the SOP implements it."),
            ("Runtime Test", "If the system performs the step, the SOP still needs trigger, checkpoint, success, failure, retry, and escalation treatment."),
        ],
    )

    doc.add_heading("How SOPs Link To Policy", level=1)
    add_matrix_table(
        doc,
        ["Policy rule", "SOP implementation", "Evidence"],
        [
            ("No invoice without Admin approval.", "SOP-BIL-01 has Admin review and approval of the billing group before invoice generation.", "Billing approval record with Admin identifier and timestamp."),
            ("Invoice must be generated from approved billing data.", "SOP-BIL-04 generates the invoice from the approved billing group and writes invoice_id back to jobs.", "Invoice record, dispatch evidence, invoiced=true write-back."),
            ("No Pickup cannot create a billable item.", "SOP-PUP-03 records valid No Pickup reasons and blocks billable records.", "No Pickup outcome record and exception monitoring where required."),
        ],
        [2650, 4300, 2410],
    )

    doc.add_heading("Moto & Co Examples", level=1)
    add_bullets(
        doc,
        [
            "SOP-BIL-01 Month-End Billing Review governs Admin review of delivered, uninvoiced jobs before invoice generation.",
            "SOP-BIL-04 Create & Send Invoice governs generation of the invoice from the approved billing group, Admin invoice review, dispatch/download boundary, invoice_id write-back, and payment monitoring handoff.",
            "SOP-DEL-04 and SOP-DEL-05 are upstream proof and completion SOPs that make a job billing-ready.",
            "The v2.0 baseline notes that V1 billing creates downloadable EOM invoice PDFs grouped by client/month, and Admin email/payment/bank reconciliation are outside the portal runtime.",
        ],
    )

    doc.add_heading("Quality Checklist", level=1)
    add_bullets(
        doc,
        [
            "The SOP title and ID match the BOAS/SOP register.",
            "The Summary sheet names parent operational policy, parent capability policy, process module, risks, and completion standard.",
            "Steps are actor-led, ordered, and testable.",
            "Platform/runtime detail is present but kept separate from business-readable steps.",
            "Data, RACI, SLA, and exception paths are complete enough for audit and UAT.",
        ],
    )

    source_note(
        doc,
        [
            "SOP/SOP-BIL-01-MonthEndBillingReview-v1.2.xlsx",
            "SOP/SOP-BIL-04-CreateSendInvoice-v1.2.xlsx",
            "baseline/v2.0/full-source/SOP/SOP_v2_baseline_manifest.csv",
            "baseline/v2.0/MotoCo_SOP_Baseline_Addendum_v1.3.md",
        ],
    )
    return save_doc(doc, "MotoCo_SOP_Reference_Guide.docx")


def main():
    paths = [
        build_capability_guide(),
        build_policy_guide(),
        build_sop_guide(),
    ]
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
